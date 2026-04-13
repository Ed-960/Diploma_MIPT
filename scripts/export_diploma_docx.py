#!/usr/bin/env python3
"""Собирает DIPLOMNAYA_RABOTA_TEKST.txt в Word (.docx). Требуется python-docx."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    print("Установите: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def is_rule_line(s: str) -> bool:
    t = s.strip()
    return len(t) >= 20 and set(t) == {"="}


def flush_buffer(doc: Document, buf: list[str]) -> None:
    if not buf:
        return
    text = " ".join(buf).strip()
    if text:
        doc.add_paragraph(text)
    buf.clear()


def add_table_from_ascii(doc: Document, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for rowline in lines:
        t = rowline.strip()
        if not t.startswith("│") or t.startswith("├"):
            continue
        parts = [p.strip() for p in rowline.split("│")]
        parts = [p for p in parts if p]
        if len(parts) >= 2:
            rows.append(parts[:2])
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, r in enumerate(rows):
        for ci in range(ncols):
            cell_text = r[ci] if ci < len(r) else ""
            tbl.rows[ri].cells[ci].text = cell_text


def text_to_docx(src: Path, dst: Path) -> None:
    raw = src.read_text(encoding="utf-8")
    lines = raw.splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    # --- титульная часть до первого блока ===...===
    titul_end = 0
    for j, line in enumerate(lines):
        if is_rule_line(line):
            titul_end = j
            break

    for j in range(titul_end):
        line = lines[j].strip()
        if line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    i = titul_end
    buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if is_rule_line(line):
            flush_buffer(doc, buf)
            i += 1
            if i >= len(lines):
                break
            title = lines[i].strip()
            i += 1
            if i < len(lines) and is_rule_line(lines[i]):
                i += 1
            if title:
                doc.add_heading(title, level=1)
            continue

        if re.match(r"^\d+\.\d+\.\s", line.strip()):
            flush_buffer(doc, buf)
            doc.add_heading(line.strip(), level=2)
            i += 1
            continue

        if line.startswith("┌"):
            flush_buffer(doc, buf)
            tbl_lines: list[str] = []
            while i < len(lines):
                tbl_lines.append(lines[i])
                if lines[i].strip().startswith("└"):
                    i += 1
                    break
                i += 1
            add_table_from_ascii(doc, tbl_lines)
            continue

        if not line.strip():
            flush_buffer(doc, buf)
            i += 1
            continue

        buf.append(line.strip())
        i += 1

    flush_buffer(doc, buf)

    doc.save(dst)
    print(f"Записано: {dst}")


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    src = root / "DIPLOMNAYA_RABOTA_TEKST.txt"
    dst = root / "DIPLOMNAYA_RABOTA_TEKST.docx"
    if not src.is_file():
        print(f"Нет файла: {src}", file=sys.stderr)
        sys.exit(1)
    text_to_docx(src, dst)


if __name__ == "__main__":
    main()
